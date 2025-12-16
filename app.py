import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt

# --- 1. LOGIC: Date Calculator ---
def add_months(source_date, months):
    month = source_date.month - 1 + months
    year = source_date.year + month // 12
    month = month % 12 + 1
    day = min(source_date.day, [31,
        29 if year % 4 == 0 and not year % 100 == 0 or year % 400 == 0 else 28,
        31, 30, 31, 30, 31, 31, 30, 31, 30, 31][month - 1])
    return date(year, month, day)

# --- 2. DATABASE: Arguments & Citations ---
KNOWLEDGE_BASE = {
    "10a1": {
        "title": "Corruption, Fraud, or Undue Means",
        "section": "9 U.S.C. § 10(a)(1)",
        "argument_header": "The Award Was Procured by Corruption, Fraud, or Undue Means.",
        "case_cite": "Bonar v. Dean Witter Reynolds, Inc., 835 F.2d 1378 (11th Cir. 1988)",
        "standard": "Movant must establish by clear and convincing evidence that the fraud was not discoverable upon the exercise of due diligence prior to or during the arbitration."
    },
    "10a2": {
        "title": "Evident Partiality",
        "section": "9 U.S.C. § 10(a)(2)",
        "argument_header": "There Was Evident Partiality in the Arbitrators.",
        "case_cite": "Commonwealth Coatings Corp. v. Continental Cas. Co., 393 U.S. 145 (1968)",
        "standard": "Arbitrators must disclose to the parties any dealings that might create an impression of possible bias."
    },
    "10a3": {
        "title": "Misconduct / Refusal to Hear",
        "section": "9 U.S.C. § 10(a)(3)",
        "argument_header": "The Arbitrators Were Guilty of Misconduct.",
        "case_cite": "Tempo Shain Corp. v. Bertek, Inc., 120 F.3d 16 (2d Cir. 1997)",
        "standard": "The panel's refusal to hear pertinent and material evidence, or to postpone the hearing, deprived Movant of a fundamentally fair hearing."
    },
    "10a4": {
        "title": "Exceeded Powers",
        "section": "9 U.S.C. § 10(a)(4)",
        "argument_header": "The Arbitrators Exceeded Their Powers.",
        "case_cite": "Oxford Health Plans LLC v. Sutter, 569 U.S. 564 (2013)",
        "standard": "The arbitrator acted outside the scope of his authority and failed to arguably construe or apply the contract."
    }
}

# --- 3. LOGIC: Document Generator ---
def generate_doc(selected_codes, filing_date):
    doc = Document()
    
    # Title / Caption
    doc.add_heading('MOTION TO VACATE ARBITRATION AWARD', 0)
    doc.add_paragraph('[INSERT CASE CAPTION HERE]')
    doc.add_paragraph(f'Date: {date.today().strftime("%B %d, %Y")}')
    
    # Introduction
    doc.add_heading('I. INTRODUCTION', level=1)
    doc.add_paragraph(
        "Movant hereby moves this Court to vacate the arbitration award dated "
        f"{filing_date.strftime('%B %d, %Y')} pursuant to the Federal Arbitration Act, 9 U.S.C. § 10."
    )
    
    # Arguments
    doc.add_heading('II. ARGUMENT', level=1)
    
    if not selected_codes:
        doc.add_paragraph("[No specific grounds selected in analysis tool.]")
    
    for code in selected_codes:
        data = KNOWLEDGE_BASE[code]
        # Add Header
        doc.add_heading(data['argument_header'], level=2)
        # Add Standard Text
        p = doc.add_paragraph()
        p.add_run("Under ").bold = True
        p.add_run(f"{data['section']}, a district court may vacate an award where {data['title'].lower()}. ")
        p.add_run("The governing standard requires a showing that: ")
        p.add_run(data['standard'])
        # Add Citation
        doc.add_paragraph(f"See generally {data['case_cite']}.", style='Intense Quote')
        # Placeholder for user facts
        doc.add_paragraph("[INSERT FACTS SPECIFIC TO THIS GROUND HERE]")

    # Conclusion
    doc.add_heading('III. CONCLUSION', level=1)
    doc.add_paragraph(
        "For the foregoing reasons, Movant respectfully requests that this Court vacate the arbitration award."
    )
    
    # Save to memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. MAIN APP ---
def main():
    st.set_page_config(page_title="FAA Vacatur Drafter", layout="wide")
    st.title("⚖️ FAA Vacatur: Analysis & Drafting Tool")

    # --- INPUTS ---
    col1, col2 = st.columns([1, 2])
    with col1:
        st.subheader("1. Timeline")
        d_award = st.date_input("Date Award Issued", value=date.today())
        deadline = add_months(d_award, 3)
        st.caption(f"Filing Deadline: {deadline.strftime('%m/%d/%Y')}")

    with col2:
        st.subheader("2. Select Grounds")
        selected_grounds = []
        c1, c2 = st.columns(2)
        with c1:
            if st.checkbox("§ 10(a)(1) Fraud"): selected_grounds.append("10a1")
            if st.checkbox("§ 10(a)(2) Partiality"): selected_grounds.append("10a2")
        with c2:
            if st.checkbox("§ 10(a)(3) Misconduct"): selected_grounds.append("10a3")
            if st.checkbox("§ 10(a)(4) Powers"): selected_grounds.append("10a4")

    st.divider()

    # --- PREVIEW SECTION ---
    if selected_grounds:
        st.subheader("3. Draft Preview")
        st.info("The following arguments will be generated in your brief:")
        for code in selected_grounds:
            st.markdown(f"**Argument:** {KNOWLEDGE_BASE[code]['argument_header']}")
            st.caption(f"Authority: *{KNOWLEDGE_BASE[code]['case_cite']}*")
    
    st.divider()

    # --- DOWNLOAD BUTTON ---
    # We generate the doc inside the button logic to ensure it captures latest state
    if selected_grounds:
        doc_file = generate_doc(selected_grounds, d_award)
        
        st.download_button(
            label="📄 Download Draft Motion (.docx)",
            data=doc_file,
            file_name="Motion_to_Vacate_Draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.warning("Please select at least one ground to enable document generation.")

if __name__ == "__main__":
    main()